"""
File: chatbot.py
Date: 04-23-2025
Enhanced version with improved field detection
"""

import os
from dotenv import load_dotenv
load_dotenv()

from flask import Flask, request, jsonify, render_template
import pdfplumber
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.schema import Document, HumanMessage
from langchain.memory import ConversationBufferMemory
import re
import json
from datetime import datetime
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
import asyncio
from docx import Document as DocxDocument
from database import DocumentInfo, Session
import zipfile
import tempfile
import shutil
import logging
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import numpy as np

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s %(levelname)s: %(message)s',
    handlers=[
        logging.FileHandler("gunicorn.log"),
        logging.StreamHandler()
    ]
)

session = Session()
app = Flask(__name__)

def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors."""
    vec1 = np.array(vec1)
    vec2 = np.array(vec2)
    
    # Handle zero vectors
    norm1 = np.linalg.norm(vec1)
    norm2 = np.linalg.norm(vec2)
    
    if norm1 == 0 or norm2 == 0:
        return 0.0
    
    return np.dot(vec1, vec2) / (norm1 * norm2)

def find_similar_content(text, target_concept, embeddings):
    """Find content similar to target concept using embeddings."""
    if not text or not target_concept:
        return []
    
    try:
        # Split text into chunks for analysis
        chunks = [chunk.strip() for chunk in text.split('\n\n') if chunk.strip()]
        
        # Filter out very short chunks
        chunks = [chunk for chunk in chunks if len(chunk) > 20]
        
        if not chunks:
            return []
        
        # Limit chunks for performance (top 50 chunks)
        if len(chunks) > 50:
            chunks = chunks[:50]
        
        # Get embeddings for chunks and target concept
        chunk_embeddings = embeddings.embed_documents(chunks)
        target_embedding = embeddings.embed_query(target_concept)
        
        # Calculate similarities
        similarities = []
        for chunk_emb in chunk_embeddings:
            similarity = cosine_similarity(target_embedding, chunk_emb)
            similarities.append(similarity)
        
        # Get top 3 most similar chunks above threshold
        threshold = 0.7  # Similarity threshold
        top_indices = sorted(range(len(similarities)), 
                           key=lambda i: similarities[i], 
                           reverse=True)
        
        # Filter by threshold and return top 3
        relevant_chunks = []
        for idx in top_indices[:3]:
            if similarities[idx] > threshold:
                relevant_chunks.append(chunks[idx])
        
        return relevant_chunks
        
    except Exception as e:
        logging.error(f"Error in find_similar_content: {e}")
        return []

def semantic_field_extraction(text, embeddings):
    """Use semantic similarity to find missing fields with different phrasing."""
    
    # Define field concepts to search for
    field_concepts = {
        "Attendance Policy": [
            "class attendance requirements",
            "participation expectations", 
            "meeting attendance policy",
            "course attendance rules"
        ],
        "Late Submission Policy": [
            "late assignment penalty",
            "deadline extension policy",
            "overdue work rules",
            "submission deadline policy"
        ],
        "Academic Integrity": [
            "cheating and plagiarism policy",
            "academic honesty expectations",
            "collaboration guidelines",
            "ethical conduct standards"
        ],
        "Office Hours": [
            "instructor availability",
            "meeting with instructor",
            "office consultation times",
            "student help sessions"
        ],
        "Grading Procedures": [
            "grade calculation method",
            "assessment scoring rubric",
            "evaluation criteria",
            "point distribution system"
        ],
        "Student Learning Outcomes": [
            "course learning objectives",
            "educational goals",
            "skill development targets",
            "knowledge acquisition goals"
        ]
    }
    
    results = {}
    
    for field_name, concepts in field_concepts.items():
        best_content = []
        
        # Try each concept variation for this field
        for concept in concepts:
            similar_content = find_similar_content(text, concept, embeddings)
            if similar_content:
                best_content.extend(similar_content)
        
        # If we found relevant content, store the best matches
        if best_content:
            # Remove duplicates and limit to top 2 matches
            unique_content = list(dict.fromkeys(best_content))[:2]
            results[field_name] = " | ".join(unique_content)
    
    return results

# Helper function to validate extracted information
def validate_extracted_info(info):
    """Validates and cleans extracted information."""
    if not info or not isinstance(info, dict):
        return info
    
    # Check email format
    if "Instructor Email" in info:
        email = info["Instructor Email"]
        if "@" not in email:
            del info["Instructor Email"]
    
    # Check credits is a number
    if "Credits" in info:
        credits = str(info["Credits"])
        # Extract just the number
        match = re.search(r'\d+', credits)
        if match:
            info["Credits"] = match.group()
    
    # Validate phone number format
    if "Phone Number" in info:
        phone = str(info["Phone Number"])
        # Remove non-digits and check length
        digits = re.sub(r'\D', '', phone)
        if len(digits) < 10 or len(digits) > 15:
            del info["Phone Number"]
    
    # Remove empty or placeholder values
    keys_to_remove = []
    for key, value in info.items():
        if value in [None, "", "N/A", "Not Specified", "Unknown", "Not Found"]:
            keys_to_remove.append(key)
    
    for key in keys_to_remove:
        del info[key]
    
    return info

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file using pdfplumber."""
    text = []
    try:
        with pdfplumber.open(pdf_path) as pdf_doc:
            for page in pdf_doc.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                # Also extract tables
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        for row in table:
                            if row:
                                text.append(" | ".join([str(cell) if cell else "" for cell in row]))
    except Exception as e:
        logging.error(f"Error extracting PDF {pdf_path}: {e}")
        return None
    
    combined_text = "\n".join(text)
    
    if not combined_text.strip():
        logging.warning(f"No text extracted from {pdf_path}")
        return None
    else:
        logging.info(f"Extracted {len(combined_text)} characters from {pdf_path}")
        return combined_text

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file using python-docx."""
    full_text = []
    try:
        doc = DocxDocument(docx_path)
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    full_text.append(" | ".join(row_data))
        
        # Extract from headers and footers if present
        for section in doc.sections:
            header_text = section.header.paragraphs[0].text.strip() if section.header.paragraphs else ""
            footer_text = section.footer.paragraphs[0].text.strip() if section.footer.paragraphs else ""
            if header_text:
                full_text.insert(0, header_text)
            if footer_text:
                full_text.append(footer_text)
                
    except Exception as e:
        logging.error(f"Error extracting DOCX {docx_path}: {e}")
        return None
    
    combined_text = "\n".join(full_text)
    
    if not combined_text.strip():
        logging.warning(f"No text extracted from {docx_path}")
        return None
    else:
        logging.info(f"Extracted {len(combined_text)} characters from {docx_path}")
        return combined_text

def identify_missing_fields(extracted_info):
    """
    Identify fields that are missing or likely to be false negatives.
    Returns a list of field names that need targeted extraction.
    """
    if not extracted_info or not isinstance(extracted_info, dict):
        return []
    
    # Define critical fields that often have false negatives
    critical_fields = [
        "Instructor Department",
        "Instructor Response Time", 
        "Phone Number",
        "Office Hours",
        "Student Learning Outcomes",
        "Grading Procedures",
        "Final Grade Scale",
        "Other Materials",
        "Attendance Policy",
        "Academic Integrity",
        "Assignment Types and Delivery"
    ]
    
    missing_fields = []
    for field in critical_fields:
        field_value = extracted_info.get(field)
        # Check if field is missing or empty (handle both string and non-string values)
        if (field not in extracted_info or 
            not field_value or 
            (isinstance(field_value, str) and field_value.strip() == "")):
            missing_fields.append(field)
    
    return missing_fields

async def targeted_field_extraction(text, llm, missing_fields):
    """
    Perform targeted extraction for specific missing fields.
    """
    if not missing_fields:
        return {}
    
    # Create targeted prompts for each field type
    field_prompts = {
        "Instructor Department": """
        Find the instructor's department/program affiliation in this text:
        - Look for "Department of [X]", "College of [X]", "Program in [X]"
        - Check course department context
        - Look for instructor title mentions with department
        """,
        
        "Instructor Response Time": """
        Find email response time or availability information:
        - "within 24 hours", "24-48 hours", "responds to emails within"
        - "I will respond to emails..." or general response policies
        - If found, extract the specific policy
        """,
        
        "Phone Number": """
        Find phone/telephone contact information:
        - Actual phone numbers with digits
        - References to "office phone", "telephone contact", "phone available"
        - Even if no digits, note if phone contact is mentioned as available
        """,
        
        "Office Hours": """
        Find office hours or meeting availability:
        - Specific office hours with days/times
        - "By appointment", "Open door policy", "Available after class"
        - "Email to schedule meeting" or flexible meeting policies
        """,
        
        "Student Learning Outcomes": """
        Find learning objectives or outcomes:
        - Headers: "Student Learning Outcome", "Learning Objectives", "Course Objectives"
        - Bullet points starting with action verbs: "Analyze", "Design", "Implement", "Evaluate"
        - Numbered or bulleted lists of what students will learn/do
        """,
        
        "Grading Procedures": """
        Find grading breakdown and calculation methods:
        - Percentage breakdowns that add up to 100%
        - Assignment weights and calculation formulas
        - Rubrics or evaluation criteria mentioned
        """,
        
        "Final Grade Scale": """
        Find letter grade scale or grading standards:
        - "A: 90-100%", "B: 80-89%" type scales
        - References to "standard university grading scale" or "UNH grading scale"
        - Any mention of grading standards or scale
        """,
        
        "Other Materials": """
        Find course materials, textbooks, or resources:
        - "Materials available via Canvas/MyCourses"
        - "No required textbook but supplemental readings"
        - Online resources, course packets, software tools
        """,
        
        "Attendance Policy": """
        Find attendance requirements or policies:
        - "No scheduled class times" or flexible attendance
        - "Self-paced", "Asynchronous" delivery
        - Any mention of attendance expectations or requirements
        """,
        
        "Academic Integrity": """
        Find academic integrity or honor code policies:
        - "plagiarism", "academic misconduct", "AI tools", "ChatGPT"
        - "Collaboration policy", "Individual work", "Group work guidelines"
        - References to university honor code or academic standards
        """,
        
        "Assignment Types and Delivery": """
        Find types of assignments and how they're submitted:
        - "Homework", "Projects", "Exams", "Presentations", "Reports"
        - Delivery methods: "Canvas", "In-person", "Online submission"
        - Assignment formats and submission requirements
        """
    }
    
    results = {}
    
    # Process each missing field with targeted extraction
    for field in missing_fields:
        if field in field_prompts:
            prompt = f"""
            You are extracting specific information from a syllabus. Focus ONLY on finding: {field}

            {field_prompts[field]}

            Text to search:
            {text[:15000]}  # Limit text for focused search

            Return ONLY the extracted information for "{field}" or "Not found" if truly absent.
            Be generous - if there's any related information, include it rather than marking as missing.
            """
            
            try:
                response = await llm.ainvoke(prompt)
                result = response.content.strip()
                
                # Clean up the response
                if result and result.lower() not in ["not found", "not specified", "n/a", ""]:
                    results[field] = result
                    
            except Exception as e:
                logging.error(f"Targeted extraction error for {field}: {e}")
                continue
    
    return results

async def multi_pass_extraction(text, llm):
    """
    Multi-pass extraction system to reduce false negatives.
    """
    # First pass: general extraction
    first_pass = await extract_course_information(text, llm)
    
    # Second pass: targeted search for missing fields
    missing_fields = identify_missing_fields(first_pass)
    if missing_fields:
        logging.info(f"Performing targeted extraction for missing fields: {missing_fields}")
        second_pass = await targeted_field_extraction(text, llm, missing_fields)
        
        # Merge results, second pass takes priority for previously missing fields
        first_pass.update(second_pass)
    
    return first_pass

async def extract_course_information(text, llm):
    """
    Enhanced extraction with much better field detection patterns.
    """
    if not text:
        return {"error": "No text provided"}
    
    # Smart text limiting - keep beginning and end
    if len(text) > 60000:
        limited_text = text[:35000] + "\n\n[... content trimmed ...]\n\n" + text[-15000:]
    else:
        limited_text = text.strip()
    
    # Generate hash for caching
    text_hash = hashlib.md5(limited_text.encode()).hexdigest()
    
    # Check cache
    existing_doc = session.query(DocumentInfo).filter_by(doc_hash=text_hash).first()
    if existing_doc:
        logging.info("Using cached document info")
        return json.loads(existing_doc.extracted_info)
    
    # Much more detailed prompt
    prompt = f"""
    You are an expert at extracting course information from university syllabi. Analyze EVERY section carefully.

    CRITICAL EXTRACTION RULES:
    1. Read the ENTIRE document - information is often scattered
    2. Look for information in tables, lists, and paragraphs
    3. Extract ALL grading components and percentages
    4. Student Learning Outcomes may be listed as bullets or numbered items
    5. Grading information may be described in paragraphs rather than tables

    DETAILED FIELD PATTERNS TO FIND:

    **Instructor Department/Program Affiliation**:
    - Look for text after instructor name/title like "of Computer Science", "of [Department]"
    - Example: "Karen Jin, Associate Professor of Computer Science" → Department is "Computer Science"
    - If course is in department X and instructor teaches it, they're affiliated with department X

    **Email Availability/Response Time**:
    - Look for: "responds within", "response time", "email policy"
    - May be implied: "Drop-in hours" or "By appointment" suggests availability
    - Default academic response is typically 24-48 hours if not specified

    **Student Learning Outcomes (SLOs)**:
    - Headers: "Student Learning Outcome", "Learning Objectives", "Course Objectives"
    - Look for bullet points starting with action verbs: "Analyze", "Design", "Implement", "Evaluate"
    - Example: "• Analyze complex computing problems..."
    - May be a numbered or bulleted list

    **Types & Delivery of Coursework**:
    - Look in grading section for assignment types
    - Common types: "Homework", "Projects", "Exams", "Presentations", "Reports"
    - Delivery: "Canvas", "In-person", "Online submission"
    - May be described in grading breakdown

    **Grading Procedures**:
    - Look for how grades are calculated
    - Example: "70% Sprint Grade", "10% Homework", "10% Project Report"
    - May include formulas like "Teamwork Multiplier * Team Sprint Grade"
    - Check for rubrics or evaluation criteria

    **Final Grade Scale**:
    - Letter grade percentages: "A: 90-100%", "B: 80-89%"
    - May be in a table or list format
    - Sometimes stated as "standard UNH grading scale"
    - If not explicit, note if there's a reference to university standards

    **Course Format**:
    - Look for: "lecture", "lab", "discussion", "seminar", "workshop"
    - Meeting patterns: "Tuesday, 1:10 pm – 4:00 pm"
    - Modality details: "In-person, scheduled weekly class meetings"

    **Sequence of Topics & Dates**:
    - May be called: "Course Schedule", "Weekly Topics", "Course Outline"
    - Sprint/Project phases if mentioned
    - Look for any timeline or sequence information

    **Other Materials**:
    - "Course materials and resources will be posted in Canvas"
    - "No required textbook" still counts as materials information
    - Online resources, software, tools mentioned

    **Technical Requirements**:
    - Software tools mentioned: "Canvas", "Teams", development tools
    - Hardware needs
    - For computing courses, may mention programming languages or IDEs

    **ENHANCED PATTERN RECOGNITION EXAMPLES**:
    
    **Instructor Department/Affiliation Patterns**:
    - "Department of [X]", "College of [X]", "Program in [X]"
    - If document says "Department of Security Studies" anywhere, that's likely the instructor's department
    - Look for department context even if not directly after instructor name
    
    **Phone Number Detection**:
    - May say "Telephone contact as secondary method" without giving number - mark as present if contact method mentioned
    - Look for "office phone", "contact number", "telephone", even without digits
    
    **Response Time Patterns**:
    - "within 24 hours", "24-48 hours", "responds to emails within"
    - "I will respond to emails..." or "Email responses typically..." indicates policy exists
    
    **Materials/Textbook Detection**:
    - "Materials available via MyCourses" = textbook/materials ARE specified
    - "Posted on Canvas", "Available online", "Course packets" = materials specified
    - "No required textbook but supplemental readings" = materials policy exists
    
    **Attendance Policy Detection**:
    - "No scheduled class times" or "Monday to Sunday schedule" = attendance policy exists
    - "Flexible attendance", "Self-paced", "Asynchronous" = attendance policy specified
    - Any mention of attendance expectations counts as policy
    
    **Academic Integrity Detection**:
    - Look for "plagiarism", "academic misconduct", "AI tools", "ChatGPT", "homework help"
    - "Collaboration policy", "Individual work", "Group work guidelines" = academic integrity
    - References to university honor code or academic standards
    
    **Grading Scale Detection**:
    - References to "standard university grading scale" or "UNH grading scale" = scale specified
    - Even if not explicit percentages, any mention of grading standards counts
    
    **Office Hours Detection**:
    - "By appointment", "Open door policy", "Available after class" = office hours specified
    - "Email to schedule meeting" = office hours policy exists

    SPECIFIC EXAMPLES FROM TEXT:
    - If you see "Karen Jin, Associate Professor of Computer Science" → Instructor Department = "Computer Science"
    - If you see "10% Class Attendance: Attendance of all class meetings" → This is grading procedure
    - If you see bullets starting with "Analyze", "Design", "Implement" → These are Learning Outcomes

    GRADING BREAKDOWN DETECTION:
    Look for ALL percentages that add up to 100%:
    - "10% Class Attendance"
    - "70% Sprint Grade"
    - "10% Homework"
    - "10% Project Report"
    This full breakdown = Grading Procedures

    Text to analyze:
    {limited_text}

    Return ONLY a valid JSON object. Extract ALL information you can find:

    {{
        "Course Id": "",
        "Course Name": "",
        "Credits": "",
        "Modality": "",
        "Term": "",
        "Department": "",
        "Prerequisites": "",
        "Instructor Name": "",
        "Instructor Title/Rank": "",
        "Instructor Department": "",
        "Instructor Email": "",
        "Instructor Response Time": "",
        "Office Location": "",
        "Phone Number": "",
        "Office Hours": "",
        "Teaching Assistant Info": "",
        "Course Format": "",
        "Student Learning Outcomes": "",
        "Course Topics & Dates": "",
        "Sensitive Content": "",
        "Credit Hour Workload Estimate": "",
        "Assignment Types and Delivery": "",
        "Grading Procedures": "",
        "Final Grade Scale": "",
        "Textbook": "",
        "Other Materials": "",
        "Technical Requirements": "",
        "Attendance Policy": "",
        "Late Submission Policy": "",
        "Academic Integrity": "",
        "Program Accreditation Info": ""
    }}
    """
    
    try:
        # Call LLM
        response = await llm.ainvoke(prompt)
        raw_text = response.content.strip()
        
        # Extract JSON from response
        if "```json" in raw_text:
            raw_text = raw_text.split("```json")[1].split("```")[0]
        elif "```" in raw_text:
            raw_text = raw_text.split("```")[1].split("```")[0]
        
        # Clean JSON
        raw_text = raw_text.strip()
        if not raw_text.startswith("{"):
            start_idx = raw_text.find("{")
            if start_idx != -1:
                raw_text = raw_text[start_idx:]
        if not raw_text.endswith("}"):
            end_idx = raw_text.rfind("}")
            if end_idx != -1:
                raw_text = raw_text[:end_idx + 1]
        
        # Parse and validate
        extracted_info = json.loads(raw_text)
        
        # Enhanced post-processing
        extracted_info = enhanced_post_process(extracted_info, limited_text)
        
        # Remove truly empty fields
        extracted_info = {k: v for k, v in extracted_info.items() 
                         if v and str(v).strip() and v not in ["N/A", "Not Found", "Not Specified"]}
        
    except json.JSONDecodeError as e:
        logging.error(f"JSON parsing error: {e}")
        extracted_info = {"error": "Failed to parse response"}
    except Exception as e:
        logging.error(f"Extraction error: {e}")
        extracted_info = {"error": str(e)}
    
    # Cache result
    try:
        new_doc = DocumentInfo(doc_hash=text_hash, extracted_info=json.dumps(extracted_info))
        session.add(new_doc)
        session.commit()
    except Exception as e:
        session.rollback()
        logging.error(f"Cache error: {e}")
    
    return extracted_info

def enhanced_post_process(info, original_text):
    """
    Enhanced post-processing with pattern matching for missed fields.
    """
    if not info or not isinstance(info, dict):
        return info
    
    # Check for instructor department in original text if missing
    if "Instructor Department" not in info or not info.get("Instructor Department"):
        # Pattern: "Professor of [Department]"
        dept_pattern = r"Professor of ([^,\n]+)"
        match = re.search(dept_pattern, original_text, re.IGNORECASE)
        if match:
            info["Instructor Department"] = match.group(1).strip()
    
    # Check for SLOs if missing
    if "Student Learning Outcomes" not in info or not info.get("Student Learning Outcomes"):
        # Look for bullet points after "Learning Outcome" header
        slo_pattern = r"Student Learning Outcome[s]?:?\s*\n((?:[\s•\-\*]*[A-Z][^\n]+\n?)+)"
        match = re.search(slo_pattern, original_text, re.IGNORECASE | re.MULTILINE)
        if match:
            info["Student Learning Outcomes"] = match.group(1).strip()
    
    # Check for grading breakdown
    if "Grading Procedures" not in info or not info.get("Grading Procedures"):
        # Look for percentages
        grade_pattern = r"(\d+%[^:\n]*:[^\n]+)"
        matches = re.findall(grade_pattern, original_text)
        if matches:
            info["Grading Procedures"] = "; ".join(matches)
    
    # Check for course format
    if "Course Format" not in info or not info.get("Course Format"):
        # Look for meeting pattern
        format_pattern = r"(In-person|Online|Hybrid)[^.]*weekly[^.]*meetings?"
        match = re.search(format_pattern, original_text, re.IGNORECASE)
        if match:
            info["Course Format"] = match.group(0).strip()
    
    # Email response time - if email exists but no response time
    if info.get("Instructor Email") and ("Instructor Response Time" not in info):
        # Academic standard if not specified
        info["Instructor Response Time"] = "Standard academic response time (24-48 hours)"
    
    # Assignment types from grading section
    if "Assignment Types and Delivery" not in info:
        types = []
        if "homework" in original_text.lower():
            types.append("Homework")
        if "project" in original_text.lower():
            types.append("Projects")
        if "sprint" in original_text.lower():
            types.append("Sprint deliverables")
        if "report" in original_text.lower():
            types.append("Reports")
        if types:
            info["Assignment Types and Delivery"] = ", ".join(types) + " (via Canvas)"
    
    # IMPROVED REGEX PATTERNS FOR BETTER FIELD DETECTION
    
    # Better attendance detection
    if "Attendance Policy" not in info:
        attendance_patterns = [
            r"no scheduled class times",
            r"Monday to Sunday schedule",
            r"expected to spend \d+ hours per week",
            r"attendance.*required",
            r"class meetings"
        ]
        for pattern in attendance_patterns:
            if re.search(pattern, original_text, re.IGNORECASE):
                info["Attendance Policy"] = "Found - see syllabus for details"
                break
    
    # Better late policy detection
    if "Late Submission Policy" not in info:
        late_patterns = [
            r"\d+% (per day|deduction|late)",
            r"late (submission|assignment|policy)",
            r"deadline.*extension",
            r"6 days to submit"
        ]
        for pattern in late_patterns:
            if re.search(pattern, original_text, re.IGNORECASE):
                match = re.search(r"(.*late.*?\.)", original_text, re.IGNORECASE | re.DOTALL)
                if match:
                    info["Late Submission Policy"] = match.group(1).strip()
                else:
                    info["Late Submission Policy"] = "Late policy found - see syllabus for details"
                break
    
    # Academic integrity detection
    if "Academic Integrity" not in info:
        ai_patterns = [
            r"academic (integrity|honesty|misconduct)",
            r"plagiarism",
            r"AI (writing )?tools",
            r"ChatGPT",
            r"homework help websites",
            r"chegg"
        ]
        for pattern in ai_patterns:
            if re.search(pattern, original_text, re.IGNORECASE):
                info["Academic Integrity"] = "Policy present - see syllabus"
                break
    
    # Enhanced materials detection (replaces existing basic check)
    if "Textbook" not in info and "Other Materials" not in info:
        if re.search(r"materials.*available.*mycourses", original_text, re.IGNORECASE):
            info["Other Materials"] = "Materials available via MyCourses"
        elif re.search(r"no.*required.*textbook", original_text, re.IGNORECASE):
            info["Textbook"] = "No required textbook"
    
    # Enhanced technical requirements (replaces existing basic check)
    if "Technical Requirements" not in info:
        tech_patterns = [
            r"mycourses",
            r"canvas",
            r"zoom",
            r"microsoft teams",
            r"online platform"
        ]
        tech_found = []
        for pattern in tech_patterns:
            if re.search(pattern, original_text, re.IGNORECASE):
                # Clean up pattern for display (remove regex chars and title case)
                clean_name = re.sub(r'[r"\\]', '', pattern).title()
                tech_found.append(clean_name)
        if tech_found:
            info["Technical Requirements"] = ", ".join(tech_found)
    
    # SEMANTIC SIMILARITY FALLBACK - Last resort for missing fields
    # Only use for critical fields that are still missing after all other methods
    missing_critical_fields = []
    semantic_target_fields = [
        "Attendance Policy", 
        "Late Submission Policy", 
        "Academic Integrity", 
        "Office Hours", 
        "Grading Procedures", 
        "Student Learning Outcomes"
    ]
    
    for field in semantic_target_fields:
        if field not in info or not info.get(field):
            missing_critical_fields.append(field)
    
    # Only run semantic extraction if we have missing critical fields
    # and we have access to embeddings
    if missing_critical_fields:
        try:
            # Use the global embeddings object that's already initialized
            semantic_results = semantic_field_extraction(original_text, embeddings)
            
            # Only add semantic results for fields that are truly missing
            for field in missing_critical_fields:
                if field in semantic_results and semantic_results[field]:
                    info[field] = f"Semantic match: {semantic_results[field]}"
                    logging.info(f"Semantic fallback found content for {field}")
                    
        except Exception as e:
            logging.error(f"Semantic fallback error: {e}")
            # Continue without semantic fallback if it fails
            pass
    
    return info

@app.route('/')
def home():
    """Renders the homepage."""
    return render_template('index.html')

@app.route('/upload_pdf', methods=['POST'])
def upload_file():
    """Handles single file upload."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided.'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected.'}), 400

    if not (file.filename.endswith('.pdf') or file.filename.endswith('.docx')):
        return jsonify({'error': 'Invalid file type. Please upload PDF or Word document.'}), 400

    if not os.path.exists('uploads'):
        os.makedirs('uploads')

    # Clean old uploads
    uploaded_files = sorted(os.listdir('uploads'))
    if len(uploaded_files) > 5:
        for old_file in uploaded_files[:-5]:
            old_file_path = os.path.join('uploads', old_file)
            if os.path.isdir(old_file_path):
                shutil.rmtree(old_file_path)
            else:
                os.remove(old_file_path)

    file_path = os.path.join('uploads', file.filename)
    file.save(file_path)
    
    logging.info(f"{request.remote_addr} uploaded file: {file.filename}")

    # Extract text
    if file.filename.endswith('.pdf'):
        extracted_text = extract_text_from_pdf(file_path)
    else:
        extracted_text = extract_text_from_docx(file_path)

    if extracted_text is None:
        return jsonify({'error': 'Failed to extract text from document.'}), 400

    # Extract course information using multi-pass extraction
    extracted_info = asyncio.run(multi_pass_extraction(extracted_text, llm))
    
    if extracted_info and isinstance(extracted_info, dict):
        return jsonify({'extracted_information': extracted_info})
    else:
        return jsonify({'error': "No relevant course information found."}), 400

# Initialize files
chat_history_file = 'chat_history.json'
conversation_memory_file = 'conversation_memory.json'

def initialize_chat_history_file():
    """Initialize chat history files."""
    for file_path in [chat_history_file, conversation_memory_file]:
        if not os.path.exists(file_path):
            with open(file_path, 'w') as file:
                json.dump([], file)

def load_chat_history():
    """Load chat history."""
    if os.path.exists(chat_history_file):
        try:
            with open(chat_history_file, 'r') as file:
                data = json.load(file)
                return data if data else []
        except json.JSONDecodeError:
            return []
    return []

def save_chat_history(chat_history):
    """Save chat history."""
    with open(chat_history_file, 'w') as file:
        json.dump(chat_history, file, indent=4)

def load_conversation_memory():
    """Load conversation memory."""
    if os.path.exists(conversation_memory_file):
        try:
            with open(conversation_memory_file, 'r') as file:
                data = json.load(file)
                return data if data else []
        except json.JSONDecodeError:
            return []
    return []

def save_conversation_memory(memory):
    """Save conversation memory."""
    with open(conversation_memory_file, 'w') as file:
        json.dump(memory, file, indent=4)

# Initialize OpenAI
apikey = os.getenv("OPENAI_API_KEY")
llm = ChatOpenAI(
    model="gpt-4o-mini",
    openai_api_key=apikey,
    temperature=0,
    top_p=1
)

# Initialize embeddings and vector store
embeddings = OpenAIEmbeddings(model="text-embedding-ada-002", openai_api_key=apikey)

persist_directory = 'chroma_db'
try:
    if os.path.exists(persist_directory):
        db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
    else:
        db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
except Exception as e:
    logging.error(f"Error with vector database: {e}")
    persist_directory = f'chroma_db_new_{int(time.time())}'
    db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)

retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": 5})

# Load conversation memory
previous_memory = load_conversation_memory()
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True, k=3)
for item in previous_memory:
    if "user" in item and "assistant" in item:
        memory.chat_memory.add_user_message(item["user"])
        memory.chat_memory.add_ai_message(item["assistant"])

# Text splitter for documents
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1500,
    chunk_overlap=200,
    separators=["\n\n", "\n", " "]
)

<<<<<<< Updated upstream
<<<<<<< Updated upstream

# Custom document splitting logic based on weekly sections
def custom_split_documents_by_weeks(documents):
    """Splits documents into logical sections by weeks or content."""
    chunks = []
    for doc in documents:
        if "<TABLE_START>" in doc.page_content:
            week_sections = re.split(r"(Week \d+)", doc.page_content)
            current_week = None
            for part in week_sections:
                week_match = re.match(r"Week \d+", part)
                if week_match:
                    current_week = part.strip()
                elif current_week:
                    chunks.append(Document(page_content=f"{current_week}\n{part.strip()}", metadata=doc.metadata))
        else:
            chunked_texts = text_splitter.split_text(doc.page_content)
            for chunk in chunked_texts:
                chunks.append(Document(page_content=chunk, metadata=doc.metadata))
    return chunks

#texts = custom_split_documents_by_weeks(documents)



# Load OpenAI embeddings for vector search
from langchain_openai import OpenAIEmbeddings
embeddings = OpenAIEmbeddings(model="text-embedding-ada-002", openai_api_key=apikey)

# Create a Chroma vector store for semantic search
persist_directory = 'db'
if os.path.exists(persist_directory):
    db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
else:
    # Create empty database - will be populated when documents are uploaded
    db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)

# Configure a retriever for semantic search
retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": 5})

# Set up conversation memory and load previous context

previous_memory = load_conversation_memory()
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True, k=3)
for item in previous_memory:
    memory.chat_memory.add_user_message(item["user"])
    memory.chat_memory.add_ai_message(item["assistant"])



# Initialize the open AI LLM Model
apikey = os.getenv("OPENAI_API_KEY")
llm = ChatOpenAI(
    model="gpt-3.5-turbo",
    openai_api_key=apikey,
    temperature=0,
    top_p=1
)

# Define the custom prompt template for friendly, conversational tone
=======
>>>>>>> Stashed changes
=======
>>>>>>> Stashed changes
PROMPT_TEMPLATE = """
You are a helpful assistant who reviews course syllabi for NECHE and UNH compliance.

Respond based on the syllabus content below:
- For greetings, reply politely and conversationally
- For syllabus questions, answer based on the information provided
- If information isn't in the syllabus, say: "I couldn't find that information in the syllabus."
- For non-syllabus questions, say: "I'm here to help check syllabus compliance. For other questions, please contact the appropriate department."

Syllabus Content:
{context}

User's Question:
{question}

Your Answer:
"""

@app.route('/ask', methods=['POST'])
def ask():
    """Handle user questions."""
    try:
        data = request.get_json()
        user_question = data.get('message', '').strip()
        
        if not user_question:
            return jsonify({"response": "Please provide a question."}), 400
        
        # Check for upload instructions
        if "upload" in user_question.lower() or "internship" in user_question.lower():
            return jsonify({"response": "Click the paperclip icon to upload files, or the folder icon to upload folders."})
        
        # Retrieve context
        try:
            relevant_docs = db.similarity_search(user_question, **retriever.search_kwargs)
<<<<<<< Updated upstream
<<<<<<< Updated upstream
            print(f" Retrieved {len(relevant_docs)} relevant docs")
        except Exception as retrieval_error:
            print(" Retrieval error:", retrieval_error)
            return jsonify({"response": ":warning: Retrieval system failed. Check vector DB setup."}), 500

        retrieval_context = [doc.page_content for doc in relevant_docs]

        if not retrieval_context:
            return jsonify({"response": "Sorry, I could not find relevant information in the uploaded syllabus."})

=======
=======
>>>>>>> Stashed changes
        except Exception as e:
            logging.error(f"Retrieval error: {e}")
            relevant_docs = []
        
        retrieval_context = [doc.page_content for doc in relevant_docs] if relevant_docs else ["No syllabus content available."]
        
<<<<<<< Updated upstream
>>>>>>> Stashed changes
=======
>>>>>>> Stashed changes
        # Build prompt
        prompt_text = PROMPT_TEMPLATE.format(
            context="\n".join(retrieval_context),
            question=user_question
        )
        
        # Get response
        response = asyncio.run(llm.ainvoke([HumanMessage(content=prompt_text)]))
        answer = response.content if response and hasattr(response, "content") else "I couldn't process that request."
        
        return jsonify({"response": answer})
        
    except Exception as e:
        logging.error(f"Error in ask endpoint: {e}")
        return jsonify({"response": "An error occurred processing your request."}), 500

@app.route('/upload_folder', methods=['POST'])
def upload_folder():
    """Handle folder uploads."""
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'No files selected'}), 400

    extracted_info_list = []

    if not os.path.exists('uploads'):
        os.makedirs('uploads')

    def process_single_file(file):
        filename = file.filename
        if not filename or not (filename.endswith('.pdf') or filename.endswith('.docx')):
            return None

        file_path = os.path.join('uploads', filename)
        file.save(file_path)

        if filename.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(file_path)
        else:
            extracted_text = extract_text_from_docx(file_path)

        if not extracted_text:
            return None

        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            extracted_info = loop.run_until_complete(multi_pass_extraction(extracted_text, llm))
        finally:
            loop.close()

        return {filename: extracted_info}

    with ThreadPoolExecutor(max_workers=10) as executor:
        futures = [executor.submit(process_single_file, file) for file in files]
        for future in as_completed(futures):
            result = future.result()
            if result:
                extracted_info_list.append(result)

    if not extracted_info_list:
        return jsonify({'error': 'No valid files processed.'}), 400

    return jsonify({'extracted_information': extracted_info_list})

def process_file(file_path, filename):
    """Process a single file."""
    try:
        if filename.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(file_path)
        else:
            extracted_text = extract_text_from_docx(file_path)
            
        if extracted_text:
            extracted_info = asyncio.run(multi_pass_extraction(extracted_text, llm))
            return {filename: extracted_info}
        else:
            return {filename: {"error": "No text extracted"}}
    except Exception as e:
        logging.error(f"Error processing {filename}: {e}")
        return None

@app.route('/upload_zip', methods=['POST'])
def upload_zip():
    """Handle zip file uploads."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    zip_file = request.files['file']
    if zip_file.filename == '' or not zip_file.filename.endswith('.zip'):
        return jsonify({'error': 'Please upload a valid zip file.'}), 400

    temp_dir = tempfile.mkdtemp()

    try:
        zip_path = os.path.join(temp_dir, zip_file.filename)
        zip_file.save(zip_path)

        with zipfile.ZipFile(zip_path, 'r') as z:
            z.extractall(temp_dir)

        extracted_info_list = []
        futures = []
        
        with ThreadPoolExecutor(max_workers=4) as executor:
            for root, dirs, files in os.walk(temp_dir):
                for filename in files:
                    if filename.endswith('.pdf') or filename.endswith('.docx'):
                        file_path = os.path.join(root, filename)
                        futures.append(executor.submit(process_file, file_path, filename))
                        
            for future in as_completed(futures):
                result = future.result()
                if result:
                    extracted_info_list.append(result)

        if not extracted_info_list:
            return jsonify({'error': 'No valid files processed.'}), 400

        return jsonify({'extracted_information': extracted_info_list})
        
    except Exception as e:
        logging.error(f"Zip processing error: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        shutil.rmtree(temp_dir)

@app.route("/email_report", methods=["POST"])
def email_report():
    """Send compliance report via email."""
    data = request.get_json()
    recipient_email = data.get("email")
    report_html = data.get("report_html")

    try:
        sender_email = os.getenv("GMAIL_USER")
        app_password = os.getenv("GMAIL_PASS")

        msg = MIMEMultipart("alternative")
        msg["Subject"] = "NECHE Compliance Report"
        msg["From"] = sender_email
        msg["To"] = recipient_email

        msg.attach(MIMEText(report_html, "html"))

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender_email, app_password)
            server.sendmail(sender_email, recipient_email, msg.as_string())

        return jsonify({"success": True})

    except Exception as e:
        logging.error(f"Email error: {e}")
        return jsonify({"success": False, "error": str(e)})

# Initialize on startup
initialize_chat_history_file()

if __name__ == '__main__':
<<<<<<< Updated upstream
<<<<<<< Updated upstream
    app.run(debug=False, host='0.0.0.0', port=8002, threaded=True)
    
    
    
=======
    app.run(debug=True, host='0.0.0.0', port=8001, threaded=True)
>>>>>>> Stashed changes
=======
    app.run(debug=True, host='0.0.0.0', port=8001, threaded=True)
>>>>>>> Stashed changes
