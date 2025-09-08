"""
Document Processing Module
This module handles extraction of text from various document formats (PDF, DOCX).
It provides functions to extract text and tables from documents for further processing.
"""

import logging
import pdfplumber
from docx import Document as DocxDocument


def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file using pdfplumber.
    
    Args:
        pdf_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text or None if extraction fails
    """
    text = []
    try:
        with pdfplumber.open(pdf_path) as pdf_doc:
            for page in pdf_doc.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                # Also extract tables
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        for row in table:
                            if row:
                                text.append(" | ".join([str(cell) if cell else "" for cell in row]))
    except Exception as e:
        logging.error(f"Error extracting PDF {pdf_path}: {e}")
        return None
    
    combined_text = "\n".join(text)
    
    if not combined_text.strip():
        logging.warning(f"No text extracted from {pdf_path}")
        return None
    else:
        logging.info(f"Extracted {len(combined_text)} characters from {pdf_path}")
        return combined_text


def extract_text_from_docx(docx_path):
    """
    Extracts text from a DOCX file using python-docx.
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text or None if extraction fails
    """
    full_text = []
    try:
        doc = DocxDocument(docx_path)
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    full_text.append(" | ".join(row_data))
        
        # Extract from headers and footers if present
        for section in doc.sections:
            header_text = section.header.paragraphs[0].text.strip() if section.header.paragraphs else ""
            footer_text = section.footer.paragraphs[0].text.strip() if section.footer.paragraphs else ""
            if header_text:
                full_text.insert(0, header_text)
            if footer_text:
                full_text.append(footer_text)
                
    except Exception as e:
        logging.error(f"Error extracting DOCX {docx_path}: {e}")
        return None
    
    combined_text = "\n".join(full_text)
    
    if not combined_text.strip():
        logging.warning(f"No text extracted from {docx_path}")
        return None
    else:
        logging.info(f"Extracted {len(combined_text)} characters from {docx_path}")
        return combined_text