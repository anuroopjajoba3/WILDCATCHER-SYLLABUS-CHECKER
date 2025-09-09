"""
File: chatbot.py
Date: 2025-09-09
Description:
- Backend switch (OpenAI <-> OSS/Ollama) via backend_switch.get_backend
- Robust syllabus JSON extraction (chunking, pre-highlighting, JSON repair)
- Alias-aware normalization and NECHE/UNH Minimal compliance checks (UI-friendly)
- Values stringified to avoid `[object Object]`
- Conservative heuristics to derive missing fields (Textbook, Technical Requirements, Office Location)
- Flask API: upload_pdf, upload_folder, upload_zip, ask, email_report
"""

from __future__ import annotations

import os
import re
import json
import time
import zipfile
import tempfile
import shutil
import hashlib
import asyncio
import logging
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Dict, Any

from dotenv import load_dotenv
load_dotenv()  # Load environment variables from .env

from flask import Flask, request, jsonify, render_template
import pdfplumber
from docx import Document as DocxDocument

# Prefer new splitters; fallback for older installs
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except Exception:
    from langchain.text_splitter import RecursiveCharacterTextSplitter  # type: ignore

from langchain.schema import Document, HumanMessage
from langchain.memory import ConversationBufferMemory

# Vector store (works with our embeddings from backend)
try:
    from langchain_chroma import Chroma  # modern
except Exception:
    from langchain_community.vectorstores import Chroma  # fallback

# Backend switch (adapter-free version recommended)
from backend_switch import get_backend

# Your DB layer
from database import DocumentInfo, Session

# ---------------- Logging ----------------
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s %(levelname)s: %(message)s',
    handlers=[logging.FileHandler("gunicorn.log"), logging.StreamHandler()]
)

# ---------------- CLI backend overrides ----------------
import argparse
parser = argparse.ArgumentParser(add_help=True)
parser.add_argument("--backend", choices=["openai", "oss"], help="Choose LLM backend (openai|oss)")
parser.add_argument("--llm-model", help='Override chat/llm model id (e.g., "gpt-4o-mini" or "llama3.1:8b")')
parser.add_argument("--embed-model", help='Override embedding model id (e.g., "text-embedding-3-small" or "sentence-transformers/all-MiniLM-L6-v2")')
args, _ = parser.parse_known_args()

# Apply CLI overrides to env (so backend_switch reads them)
if args.backend:
    os.environ["LLM_BACKEND"] = args.backend
if args.llm_model:
    os.environ["LLM_MODEL"] = args.llm_model
if args.embed_model:
    os.environ["EMBEDDING_MODEL"] = args.embed_model

# ---------------- Build backend ----------------
backend = get_backend(os.getenv("LLM_BACKEND"))
llm = backend.llm
embeddings = backend.embeddings
print(f"[backend] {backend.name} | llm={os.getenv('LLM_MODEL') or os.getenv('OLLAMA_CHAT_MODEL')} | embed={os.getenv('EMBEDDING_MODEL') or os.getenv('HF_EMBED_MODEL')}")

# ---------------- Flask app & DB ----------------
session = Session()
app = Flask(__name__)

# ---------------- Alias map & compliance config ----------------
FIELD_ALIASES: Dict[str, List[str]] = {
    "Textbook": [
        "Required Textbook", "Required Textbooks", "Required Readings",
        "Course Materials", "Reading Materials", "Readings", "Required Material(s)"
    ],
    "Technical Requirements": [
        "System Requirements", "Technology Requirements", "Computer Requirements",
        "Software/Hardware Requirements", "Technology/Requirements"
    ],
    "Phone Number": [
        "Contact Number", "Telephone", "Tel", "Cell", "Mobile", "Office Phone"
    ],
    "Final Grade Scale": [
        "Grading Scale", "Grade Breakdown", "Letter Grades", "Final Grading Scale",
        "Final Grades", "Final Grade Breakdown"
    ],
    "Grading Procedures": [
        "Grading Policy", "Grading Policies", "Grades", "Assessment & Grading",
        "Evaluation & Grading"
    ],
    "Modality": [
        "Time and Location", "Course Delivery", "Delivery Modality", "Meeting Modality"
    ],
}

# Canonical keys
NECHE_KEYS: List[str] = [
    "Course Id","Course Name","Credits","Modality","Term","Department","Prerequisites",
    "Instructor Name","Instructor Title/Rank","Instructor Department","Instructor Email",
    "Instructor Response Time","Office Location","Phone Number","Office Hours",
    "Teaching Assistant Info","Course Format","Student Learning Outcomes","Course Topics & Dates",
    "Sensitive Content","Credit Hour Workload Estimate","Assignment Types and Delivery",
    "Grading Procedures","Final Grade Scale","Textbook","Other Materials","Technical Requirements",
    "Attendance Policy","Late Submission Policy","Academic Integrity","Program Accreditation Info"
]

# Requirements (tune as needed)
NECHE_REQUIRED = {
    "Course Id", "Course Name", "Credits", "Modality",
    "Instructor Name", "Instructor Email", "Phone Number",
    "Grading Procedures", "Final Grade Scale", "Academic Integrity"
}
UNH_MIN_REQUIRED = {
    "Textbook", "Technical Requirements"
}

# ---- UI legacy aliasing (to match your frontend labels) ----
LEGACY_UI_ALIASES = {
    "Office Location": ["Office address"],
    "Student Learning Outcomes": ["Course SLOs/Program SLOs", "SLOs"],
    "Assignment Types and Delivery": ["Types & delivery of coursework"],
    "Course Topics & Dates": ["Sequence of course topics and important dates"],
    "Textbook": ["Required/recommended textbook"],
    "Technical Requirements": ["Technical requirements"],
    "Final Grade Scale": ["Final grade scale"],
}

def _norm(s: str) -> str:
    return (s or "").strip().lower()

def _stringify_value(v):
    """Turn any JSON-like value into a readable string for the UI."""
    if v is None:
        return ""
    if isinstance(v, str):
        return v.strip()
    if isinstance(v, (int, float, bool)):
        return str(v)
    if isinstance(v, list):
        if v and all(isinstance(x, dict) for x in v):
            return "\n".join("- " + ", ".join(f"{k}: {str(val)}" for k, val in x.items()) for x in v)
        return "\n".join(f"- {_stringify_value(x)}" for x in v)
    if isinstance(v, dict):
        return "\n".join(f"- {k}: {_stringify_value(val)}" for k, val in v.items())
    return str(v)

def coerce_values_to_str(d: dict) -> dict:
    """Return a copy with every value coerced to a clean string for UI."""
    out = {}
    for k, v in (d or {}).items():
        out[k] = _stringify_value(v)
    return out

def add_legacy_ui_keys(canonical: dict) -> dict:
    """
    Add UI-legacy duplicate keys so the frontend finds what it expects,
    while keeping the canonical keys too.
    """
    out = dict(canonical)
    for canon, legacy_names in LEGACY_UI_ALIASES.items():
        if canon in canonical and _stringify_value(canonical.get(canon)):
            for legacy in legacy_names:
                if not _stringify_value(out.get(legacy)):
                    out[legacy] = canonical[canon]
    return out

def normalize_with_aliases(data: Dict[str, Any]) -> Dict[str, Any]:
    """Move alias values into canonical keys, prefer non-empty canonical, drop leftover alias keys."""
    out = dict(data) if isinstance(data, dict) else {}
    for canonical, aliases in FIELD_ALIASES.items():
        current_val = out.get(canonical, "")
        if not isinstance(current_val, str):
            current_val = str(current_val) if current_val is not None else ""
        if current_val.strip():
            for alias in aliases:
                out.pop(alias, None)
            continue
        for alias in aliases:
            alias_val = out.get(alias, "")
            if isinstance(alias_val, (list, dict)):
                alias_val = json.dumps(alias_val, ensure_ascii=False)
            alias_val = str(alias_val).strip() if alias_val is not None else ""
            if alias_val:
                out[canonical] = alias_val
                out.pop(alias, None)
                break
        for alias in aliases:
            out.pop(alias, None)
    return out

def has_nonempty(d: Dict[str, Any], key: str) -> bool:
    """Treat 'Not Found', 'N/A', 'None' as empty; lists/dicts must be non-empty."""
    if key not in d:
        return False
    val = d.get(key)
    if val is None:
        return False
    if isinstance(val, str):
        s = val.strip()
        if s == "":
            return False
        s_low = s.lower()
        if s_low in {"not found", "n/a", "na", "none"}:
            return False
        return True
    if isinstance(val, (list, dict)):
        return len(val) > 0
    return True

def has_nonempty_with_legacy(d: dict, canonical_key: str) -> bool:
    if has_nonempty(d, canonical_key):
        return True
    for legacy in LEGACY_UI_ALIASES.get(canonical_key, []):
        if has_nonempty(d, legacy):
            return True
    return False

def evaluate_compliance_ui_friendly(d: Dict[str, Any]) -> Dict[str, Any]:
    """
    Treat legacy UI keys as valid sources for the canonical requirement.
    Also add legacy duplicates so UI can render them directly.
    """
    nd = normalize_with_aliases(d)
    nd = add_legacy_ui_keys(nd)

    neche_missing = [k for k in sorted(NECHE_REQUIRED) if not has_nonempty_with_legacy(nd, k)]
    unh_missing = [k for k in sorted(UNH_MIN_REQUIRED) if not has_nonempty_with_legacy(nd, k)]

    return {
        "neche_ok": len(neche_missing) == 0,
        "neche_missing": neche_missing,
        "unh_ok": len(unh_missing) == 0,
        "unh_missing": unh_missing,
        "normalized": nd,
    }

# ---------------- Heuristics for deriving missing fields ----------------
def _contains_any(text: str, words: List[str]) -> bool:
    t = text.lower()
    return any(w.lower() in t for w in words)

def _find_office_location(raw: str) -> str:
    """
    Try to find an office location/address line in the raw syllabus text.
    Strategies:
      1) Lines that start with 'Office', 'Office Location', 'Office Address'
      2) Common patterns with 'Hall' and 'Room'
      3) 'Office:' anywhere in line
    Returns a short, cleaned string (or "" if not found).
    """
    if not raw:
        return ""
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    # 1) Labeled lines
    label_re = re.compile(r'(?i)^(Office(?:\s+(?:Location|Address))?)\s*[:\-]\s*(.+)$')
    for ln in lines:
        m = label_re.match(ln)
        if m:
            val = m.group(2).strip()
            # cut at extra contact info if present
            val = re.split(r'(?i)\b(Phone|Tel|Email|E-mail)\b', val)[0].strip()
            return val[:160]

    # 2) Hall + Room patterns (e.g., "Kingsbury Hall Room 201", "McConnell 210")
    hall_room_re = re.compile(r'(?i)\b([A-Z][a-zA-Z]+(?:\s+Hall)?)\s*(?:Room\s*)?([A-Za-z]?\d{1,4})\b')
    for ln in lines:
        if hall_room_re.search(ln):
            # return the phrase around the match
            return ln[:160]

    # 3) Any 'Office:' mention
    inline_re = re.compile(r'(?i)\bOffice\s*[:\-]\s*([^|,;]+)')
    for ln in lines:
        m = inline_re.search(ln)
        if m:
            return m.group(1).strip()[:160]

    return ""

def derive_missing_fields_from_text(raw_text: str, extracted: dict) -> dict:
    """
    If the LLM missed some obvious fields, infer them conservatively from the raw text.
    - Textbook: LMS-provided materials
    - Technical Requirements: Online/LMS baseline
    - Office Location: labeled lines or typical building/room patterns
    """
    out = dict(extracted)
    raw = (raw_text or "").strip()

    # TEXTBOOK
    if not has_nonempty(out, "Textbook"):
        if _contains_any(raw, ["available via mycourses", "available on mycourses", "canvas", "posted on mycourses", "posted online in the course site"]):
            out["Textbook"] = "Materials provided via LMS (MyCourses/Canvas)"

    # TECHNICAL REQUIREMENTS
    if not has_nonempty(out, "Technical Requirements"):
        modality = (out.get("Modality") or "").lower()
        if "online" in modality or _contains_any(raw, ["mycourses", "canvas", "lms", "online", "internet", "web browser"]):
            out["Technical Requirements"] = "Reliable internet + access to LMS (MyCourses/Canvas) + modern web browser"

    # OFFICE LOCATION / ADDRESS
    if not has_nonempty(out, "Office Location"):
        guess = _find_office_location(raw)
        if guess:
            out["Office Location"] = guess

    return out

# ---------------- PDF / DOCX extraction ----------------
def extract_text_from_pdf(pdf_path: str):
    parts = []
    with pdfplumber.open(pdf_path) as pdf_doc:
        for page in pdf_doc.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)
    text = "\n".join(parts)
    if not text.strip():
        print(f"Warning: No text extracted from {pdf_path}")
        return None
    print(f"Extracted {len(text)} characters from {pdf_path}")
    return text

def extract_text_from_docx(docx_path: str):
    full_text = []
    doc = DocxDocument(docx_path)

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                full_text.append(" | ".join(row_data))

    combined = "\n".join(full_text)
    if not combined.strip():
        print(f"Warning: No text extracted from {docx_path}")
        return None
    print(f"Extracted {len(combined)} characters from {docx_path}")
    return combined

# ---------------- History & memory ----------------
chat_history_file = 'chat_history.json'
conversation_memory_file = 'conversation_memory.json'

def initialize_chat_history_file():
    if not os.path.exists(chat_history_file):
        with open(chat_history_file, 'w') as f:
            json.dump([], f)
    if not os.path.exists(conversation_memory_file):
        with open(conversation_memory_file, 'w') as f:
            json.dump([], f)

def load_chat_history():
    if os.path.exists(chat_history_file):
        try:
            with open(chat_history_file, 'r') as f:
                data = json.load(f)
                return data if data else []
        except json.JSONDecodeError:
            return []
    return []

def save_chat_history(chat_history):
    with open(chat_history_file, 'w') as f:
        json.dump(chat_history, f, indent=4)

def load_conversation_memory():
    if os.path.exists(conversation_memory_file):
        try:
            with open(conversation_memory_file, 'r') as f:
                data = json.load(f)
                return data if data else []
        except json.JSONDecodeError:
            return []
    return []

def save_conversation_memory(memory):
    with open(conversation_memory_file, 'w') as f:
        json.dump(memory, f, indent=4)

def get_question_hash(question: str):
    normalized = " ".join(question.lower().split())
    return hashlib.md5(normalized.encode()).hexdigest()

def get_cached_response(question_hash):
    chat_history = load_chat_history()
    for entry in chat_history:
        if entry.get("question_hash") == question_hash:
            return entry.get("bot_response")
    return None

def save_interaction_to_json(user_question, bot_response):
    chat_history = load_chat_history()
    question_hash = get_question_hash(user_question)
    new_chat = {
        "user_question": user_question,
        "bot_response": bot_response,
        "timestamp": datetime.now().isoformat(),
        "question_hash": question_hash
    }
    for entry in chat_history:
        if entry.get("question_hash") == question_hash:
            entry.update(new_chat)
            break
    else:
        chat_history.append(new_chat)
    save_chat_history(chat_history)

# ---------------- Splitters & helpers ----------------
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1500,
    chunk_overlap=200,
    separators=["\n\n", "\n", " "]
)

def custom_split_documents_by_weeks(documents: List[Document]) -> List[Document]:
    chunks: List[Document] = []
    for doc in documents:
        if "<TABLE_START>" in doc.page_content:
            week_sections = re.split(r"(Week \d+)", doc.page_content)
            current_week = None
            for part in week_sections:
                week_match = re.match(r"Week \d+", part)
                if week_match:
                    current_week = part.strip()
                elif current_week:
                    chunks.append(Document(page_content=f"{current_week}\n{part.strip()}", metadata=doc.metadata))
        else:
            chunked_texts = text_splitter.split_text(doc.page_content)
            for chunk in chunked_texts:
                chunks.append(Document(page_content=chunk, metadata=doc.metadata))
    return chunks

def extract_texts_from_multiple_pdfs(pdf_directory: str) -> List[Document]:
    documents: List[Document] = []
    for pdf_file in os.listdir(pdf_directory):
        if pdf_file.endswith(".pdf"):
            pdf_path = os.path.join(pdf_directory, pdf_file)
            pdf_text = extract_text_from_pdf(pdf_path)
            if pdf_text:
                documents.append(Document(page_content=pdf_text, metadata={"source": pdf_file}))
    return documents

# ---------------- Vector DB (Chroma) ----------------
persist_directory = 'chroma_db'
try:
    db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)
except Exception as e:
    print(f"Error loading existing database: {e}")
    print("Creating fresh database...")
    persist_directory = f'chroma_db_new_{int(time.time())}'
    db = Chroma(persist_directory=persist_directory, embedding_function=embeddings)

retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": 5})

# ---------------- Conversation memory ----------------
previous_memory = load_conversation_memory()
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True, k=3)
for item in previous_memory:
    memory.chat_memory.add_user_message(item.get("user", ""))
    memory.chat_memory.add_ai_message(item.get("assistant", ""))

# ---------------- LLM helpers ----------------
async def llm_text_async(prompt_text: str) -> str:
    """
    Try chat-style first (HumanMessage), then fall back to plain string
    (so it works for ChatOpenAI AND Ollama community LLMs).
    """
    try:
        resp = await llm.ainvoke([HumanMessage(content=prompt_text)])
    except Exception:
        resp = await llm.ainvoke(prompt_text)
    return getattr(resp, "content", str(resp))

def try_parse_json_str(s: str):
    """Try to parse JSON with light repairs for common issues."""
    raw = s.strip()
    try:
        with open("last_raw_llm_output.txt", "w", encoding="utf-8") as f:
            f.write(raw)
    except Exception:
        pass

    l = raw.find("{")
    r = raw.rfind("}")
    if l != -1 and r != -1 and r > l:
        raw = raw[l:r+1]

    raw = raw.replace("“", "\"").replace("”", "\"").replace("’", "'").replace("‘", "'")
    raw = re.sub(r",\s*([}\]])", r"\1", raw)

    try:
        return json.loads(raw), None
    except json.JSONDecodeError as e:
        return None, f"JSON parse error: {e}"

def merge_dicts_fill_missing(base: dict, add: dict) -> dict:
    out = dict(base) if base else {}
    for k, v in (add or {}).items():
        if k not in out or (isinstance(out[k], str) and not out[k].strip()):
            out[k] = v
    return out

def chunk_text(text: str, max_len: int = 20000) -> List[str]:
    if len(text) <= max_len:
        return [text]
    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(start + max_len, len(text))
        cut = text.rfind("\n", start, end)
        if cut == -1 or cut <= start + int(0.4 * max_len):
            cut = end
        chunks.append(text[start:cut])
        start = cut
    return chunks

def preprocess_text_for_extraction(text: str) -> str:
    """
    Pre-highlight patterns (phone, grade scale) so LLM won’t miss them.
    """
    # phone numbers
    text = re.sub(r'(\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b)', r'PHONE_NUMBER:\1', text)
    # grade scale lines like "A: 94% to 100%" etc.
    text = re.sub(r'(^[A-F][+-]?:\s*[^\\n]*\d+%[^\\n]*$)', r'GRADE_SCALE:\1', text, flags=re.MULTILINE)
    return text

def json_only_prompt(limited_text: str) -> str:
    return f"""
You are extracting syllabus fields. Output ONLY a valid JSON object—no extra text.

Rules:
- If a field is missing in the text, INCLUDE the key with empty string "".
- Map synonyms/aliases back to the correct key (e.g., Course Materials -> Textbook; System Requirements -> Technical Requirements).
- If you see PHONE_NUMBER:xxx, put it under "Phone Number".
- If you see GRADE_SCALE: lines, put them under "Final Grade Scale".
- Course id and course number are the same.
- "Time and Location" is the same as "Modality".
- If a phrase follows "Instructor Title/Rank", treat it as "Instructor Department".

Target keys (always include; blank if unknown):
{json.dumps(NECHE_KEYS, indent=2)}

TEXT:
{limited_text}
""".strip()

# ---------------- Robust course info extraction ----------------
async def extract_course_information(text: str, llm_instance):
    """
    - up to 60k chars straight; chunk if larger
    - pre-highlight phone/grade scale
    - strict JSON-only prompt
    - JSON repair + merge across chunks
    - cache by limited_text hash
    - derive missing fields from full raw text (Textbook, Technical Requirements, Office Location)
    """
    MAX_PASS_LEN = 60000
    processed = preprocess_text_for_extraction(text)
    limited_text = processed[:MAX_PASS_LEN].strip()
    text_hash = hashlib.md5(limited_text.encode()).hexdigest()

    existing_doc = session.query(DocumentInfo).filter_by(doc_hash=text_hash).first()
    if existing_doc:
        print("Using cached document info.")
        cached = json.loads(existing_doc.extracted_info)
        for k in NECHE_KEYS:
            cached.setdefault(k, "")
        # Derive from full raw text as well (for new heuristics), but don't re-save cache here
        try:
            cached = derive_missing_fields_from_text(text, cached)
        except Exception:
            pass
        return cached

    source_text = processed if len(processed) > MAX_PASS_LEN else limited_text
    parts = chunk_text(source_text, max_len=20000)

    merged: dict = {}
    errors: List[str] = []

    for idx, part in enumerate(parts, start=1):
        prompt = json_only_prompt(part)
        raw = await llm_text_async(prompt)
        parsed, err = try_parse_json_str(raw)
        if parsed is None:
            errors.append(f"Chunk {idx}: {err}")
            continue
        merged = merge_dicts_fill_missing(merged, parsed)

    for k in NECHE_KEYS:
        merged.setdefault(k, "")

    if not merged:
        result = {"error": "Failed to parse LLM response", "details": errors[:3]}
    else:
        result = merged

    # Derive missing fields from full raw text
    try:
        result = derive_missing_fields_from_text(text, result)
    except Exception:
        pass

    try:
        new_doc = DocumentInfo(doc_hash=text_hash, extracted_info=json.dumps(result))
        session.add(new_doc)
        session.commit()
    except Exception as e:
        session.rollback()
        print("Error inserting document info:", e)

    return result

# ---------------- Routes ----------------
@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload_pdf', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided.'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected. Please upload a valid document.'}), 400

    if not (file.filename.endswith('.pdf') or file.filename.endswith('.docx')):
        return jsonify({'error': 'Invalid file type. Please upload a PDF or Word document (.docx).'}), 400

    if not os.path.exists('uploads'):
        os.makedirs('uploads')

    # Trim uploads directory (keep last 5)
    uploaded_files = sorted(os.listdir('uploads'))
    if len(uploaded_files) > 5:
        for old_file in uploaded_files[:-5]:
            old_file_path = os.path.join('uploads', old_file)
            if os.path.isdir(old_file_path):
                shutil.rmtree(old_file_path)
            else:
                os.remove(old_file_path)

    file_path = os.path.join('uploads', file.filename)
    file.save(file_path)
    logging.info(f"{request.remote_addr} uploaded file: {file.filename}")

    extracted_text = extract_text_from_pdf(file_path) if file.filename.endswith('.pdf') else extract_text_from_docx(file_path)
    if extracted_text is None:
        return jsonify({'error': 'Failed to extract text from the document. Please check the file format.'}), 400

    extracted_info = asyncio.run(extract_course_information(extracted_text, llm))
    print("Extracted Course Information:", extracted_info)

    # UI-friendly normalization + compliance + stringification
    comp = evaluate_compliance_ui_friendly(extracted_info)
    normalized_str = coerce_values_to_str(comp["normalized"])
    normalized_str = add_legacy_ui_keys(normalized_str)

    result_payload = {
        "extracted_information": normalized_str,
        "neche": {"ok": comp["neche_ok"], "missing": comp["neche_missing"]},
        "unh_minimal": {"ok": comp["unh_ok"], "missing": comp["unh_missing"]},
    }
    return jsonify(result_payload)

@app.route('/ask', methods=['POST'])
def ask():
    try:
        data = request.get_json()
        user_question = data.get('message', '')
        if not isinstance(user_question, str) or not user_question.strip():
            return jsonify({"response": "Hmm, I didn't quite catch that. Could you try rephrasing?"}), 400

        normalized_q = user_question.strip().lower()
        if "internship" in normalized_q or "upload" in normalized_q:
            return jsonify({"response": "Click on the 📎 icon to upload file/zip and 📁 to upload a folder below. I'll check compliance for you."})

        try:
            relevant_docs = db.similarity_search(user_question, **retriever.search_kwargs)
            print(f"Retrieved {len(relevant_docs)} relevant docs")
        except Exception as retrieval_error:
            print("Retrieval error:", retrieval_error)
            return jsonify({"response": ":warning: Retrieval system failed. Check vector DB setup."}), 500

        retrieval_context = [doc.page_content for doc in relevant_docs] or ["No syllabus content available."]

        PROMPT_TEMPLATE = """
You are a helpful and friendly assistant who reviews course syllabi to determine if they meet NECHE and UNH minimal compliance requirements.

You should respond to questions based on the syllabus content below.

- If the question is a greeting or casual small talk (e.g., "hi", "how are you"), reply politely and conversationally.
- If the question is about syllabus content (like instructor name, grading, schedule), answer based on the information provided.
- If the answer is not in the syllabus, respond with: "I'm sorry, I couldn't find that information in the syllabus."
- If the question is unrelated to NECHE or syllabi, respond with: "I'm here to help check if your uploaded syllabus meets NECHE and minimal compliance requirements. For other questions, please contact the appropriate department."

---
### Syllabus Content:
{context}

### User's Question:
{question}
---
Your Answer:
""".strip()

        prompt_text = PROMPT_TEMPLATE.format(
            context="\n".join(retrieval_context),
            question=user_question.strip()
        )
        print("📄 Prompt (preview):\n", prompt_text[:1000])

        try:
            answer = asyncio.run(llm_text_async(prompt_text))
            if not answer:
                raise ValueError("No content returned from LLM.")
            print("LLM Response:", answer[:500])
        except Exception as llm_error:
            print("LLM Error:", llm_error)
            return jsonify({"response": f":warning: LLM error: {str(llm_error)}"}), 500

        return jsonify({"response": answer})

    except Exception as e:
        print("Server error:", e)
        return jsonify({"response": f":warning: Server error: {str(e)}"}), 500

@app.route('/upload_folder', methods=['POST'])
def upload_folder():
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'No files selected'}), 400

    extracted_info_list = []
    if not os.path.exists('uploads'):
        os.makedirs('uploads')

    def process_single_file(file):
        filename = file.filename
        if filename == '' or not (filename.endswith('.pdf') or filename.endswith('.docx')):
            return None

        file_path = os.path.join('uploads', filename)
        file.save(file_path)

        extracted_text = extract_text_from_pdf(file_path) if filename.endswith('.pdf') else extract_text_from_docx(file_path)
        if not extracted_text:
            return None

        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            extracted_info = loop.run_until_complete(extract_course_information(extracted_text, llm))
        finally:
            loop.close()

        comp = evaluate_compliance_ui_friendly(extracted_info)
        normalized_str = coerce_values_to_str(comp["normalized"])
        normalized_str = add_legacy_ui_keys(normalized_str)

        return {
            filename: {
                "extracted_information": normalized_str,
                "neche": {"ok": comp["neche_ok"], "missing": comp["neche_missing"]},
                "unh_minimal": {"ok": comp["unh_ok"], "missing": comp["unh_missing"]},
            }
        }

    with ThreadPoolExecutor(max_workers=10) as executor:
        futures = [executor.submit(process_single_file, file) for file in files]
        for future in as_completed(futures):
            result = future.result()
            if result:
                extracted_info_list.append(result)

    if not extracted_info_list:
        return jsonify({'error': 'No valid files processed.'}), 400

    return jsonify({'extracted_information': extracted_info_list})

def process_file(file_path, filename):
    try:
        extracted_text = extract_text_from_pdf(file_path) if filename.endswith('.pdf') else extract_text_from_docx(file_path)
        if extracted_text:
            extracted_info = asyncio.run(extract_course_information(extracted_text, llm))
            comp = evaluate_compliance_ui_friendly(extracted_info)
            normalized_str = coerce_values_to_str(comp["normalized"])
            normalized_str = add_legacy_ui_keys(normalized_str)
            return {
                filename: {
                    "extracted_information": normalized_str,
                    "neche": {"ok": comp["neche_ok"], "missing": comp["neche_missing"]},
                    "unh_minimal": {"ok": comp["unh_ok"], "missing": comp["unh_missing"]},
                }
            }
        else:
            return {filename: {"error": "No text extracted"}}
    except Exception as e:
        print(f"Error processing {filename}: {e}")
    return None

@app.route('/upload_zip', methods=['POST'])
def upload_zip():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    zip_file = request.files['file']
    if zip_file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    if not zip_file.filename.endswith('.zip'):
        return jsonify({'error': 'Invalid file type. Please upload a zip file.'}), 400

    temp_dir = tempfile.mkdtemp()

    try:
        zip_path = os.path.join(temp_dir, zip_file.filename)
        zip_file.save(zip_path)

        with zipfile.ZipFile(zip_path, 'r') as z:
            z.extractall(temp_dir)

        extracted_info_list = []
        futures = []
        with ThreadPoolExecutor(max_workers=4) as executor:
            for root, dirs, files in os.walk(temp_dir):
                for filename in files:
                    if filename.endswith('.pdf') or filename.endswith('.docx'):
                        file_path = os.path.join(root, filename)
                        futures.append(executor.submit(process_file, file_path, filename))
            for future in as_completed(futures):
                result = future.result()
                if result:
                    extracted_info_list.append(result)

        if not extracted_info_list:
            return jsonify({'error': 'No valid files processed.'}), 400

        return jsonify({'extracted_information': extracted_info_list})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        shutil.rmtree(temp_dir)

# ---------------- Email report ----------------
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

@app.route("/email_report", methods=["POST"])
def email_report():
    data = request.get_json()
    recipient_email = data.get("email")
    report_html = data.get("report_html")

    try:
        sender_email = os.getenv("GMAIL_USER")
        app_password = os.getenv("GMAIL_PASS")

        msg = MIMEMultipart("alternative")
        msg["Subject"] = "NECHE Compliance Report"
        msg["From"] = sender_email
        msg["To"] = recipient_email
        msg.attach(MIMEText(report_html, "html"))

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender_email, app_password)
            server.sendmail(sender_email, recipient_email, msg.as_string())

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

# ---------------- Boot ----------------
initialize_chat_history_file()

if __name__ == '__main__':
    # With your .env, run:
    #   python chatbot.py --backend oss
    # Ensure: ollama pull llama3.1:8b   (once)
    app.run(debug=False, host='0.0.0.0', port=8001, threaded=True)
